from docx import Document
from docx.shared import Inches,Mm
from docx.enum.style import WD_STYLE_TYPE
from pandas import DataFrame as DF
from pandas import *
import pandas as pd
import matplotlib.pyplot as plt
#Adding table to word doc function
def Add_Data_Frame_to_Word(DataFrame,doc_name):
    df=DataFrame  # Assigning the Dataframe to df Variable
    doc_table = doc_name.add_table(df.shape[0]+1, df.shape[1]+1)
    
    # add the header rows.
    for j in range(df.shape[-1]):
        #print(df.shape)
        if j==0: 
            doc_table.cell(0,j).text = df.columns.name
            doc_table.cell(0,j+1).text = 'Bin'+str(df.columns[j])
        else: 
            doc_table.cell(0,j+1).text = 'Bin'+str(df.columns[j])

    # add the rest of the data frame
    for i in range(df.shape[0]):
        doc_table.cell(i+1,0).text = str(df.index[i])
        for j in range(df.shape[-1]):
            doc_table.cell(i+1,j+1).text = str(df.values[i,j])
    doc_table.style = 'TableGrid'
    doc_name.add_paragraph()

data=pd.read_csv('D:\\2360\\Scripting_data\\Python_trail\\GIT_HUB\\Tool\\STDF-READER-\\Input_to_clear\\Raw data\\W649X219_RT1_C633Y5F0301_7DBS5V_TS-HP9-014_20161206014854.csv')
dataframe=DF(data)

document = Document()

def Add_Header_Fotter(self,doc_name):
    document=doc_name
    section=document.sections[0]
    #section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    ## Adding Header and Fotter
    header=section.header
    header_table =header.add_table(rows=3, cols=4,width= Inches(9.0))
    header_table.cell(0,1).merge(header_table.cell(0,3))
    header_table.cell(0,0).merge(header_table.cell(2,0))
    header_table.cell(0,0).width=1.0
    #header_table.autofit==True
    header_table.cell(0,1).text='GOURISANKAR_TOOL'

    header_table.cell(1,1).text='Title'
    header_table.cell(1,1).width=Inches(0.01)
    header_table.cell(1,2).text='Test Result Report - Finger Print'
    header_table.cell(1,2).width=4.0
    header_table.cell(2,1).text='Doc No'
    header_table.cell(2,2).text='TBD'
    header_table.style = 'TableGrid'
    try:
        header_table.rows[0].cells[0].add_paragraph(0).add_run().add_picture('D:/2360/Scripting_data/Python_trail/GIT_HUB/Tool/STDF-READER-/Tool_GUI/Icons/Logo.png',width=Inches(0.8))
    except:
        pass
    fotter=section.footer
    fotter.paragraphs[0].text='GOURISANKAR_TOOL'
    # Header and Fotter completed

# Begining the documentation

side_heading=document.add_heading('Introduction:', level=2)
side_heading.bold = True
document.add_paragraph('\t\t This report consists of yield and failure analysis for the lot DS5_DS5_T8C342-06F2_CP1_TS-HP9-014_DS5_CP  consist of 298 devices of wafer sort. ')

#Single header
#HardBin Summary

'''side_heading=document.add_heading('Over All Hardware Binning:', level=2)
side_heading.bold = True'''

def Binning_and_plotting(Dataframe_name,value_field,index_column,Name_of_binning,columns_column=None):
    ## Adding Hbin Summary to the Document
    Bin_summary=pd.pivot_table(dataframe,values=value_field,index=[index_column],columns=columns_column,aggfunc=len,fill_value=0)
    side_heading=document.add_heading('Over All %s Binning:'%Name_of_binning, level=2)
    side_heading.bold = True
    if columns_column==None:
        Bin_summary=Bin_summary.sort_values(by=value_field,ascending=False)
        ax=Bin_summary.plot(kind='bar',x=Bin_summary.index,y=value_field,stacked=False,grid=False)
    elif columns_column!=None:
        Bin_summary['Total']= Bin_summary.sum(numeric_only=True, axis=1)
        Bin_summary=Bin_summary.sort_values(by='Total',ascending=False)
        del Bin_summary['Total']
        ax=Bin_summary.plot(kind='bar',x=Bin_summary.index,stacked=False,grid=False)
    plt.ylabel('Count')
    plt.title(Name_of_binning+' Bin Pareto')
    plt.tight_layout()
    plt.savefig(Name_of_binning+'.png')
    document.add_picture(Name_of_binning+'.png')
    document.add_paragraph()
    if columns_column==None:        
        Bin_summary.loc['Total']= Bin_summary.sum(numeric_only=True, axis=0)
        Bin_summary['%'] = Bin_summary.apply(lambda row: "{0:.2f}%".format((row[value_field]/max(Bin_summary[value_field]))*100), axis = 1)
    elif columns_column!=None: 
        Bin_summary.loc['Total']= Bin_summary.sum(numeric_only=True, axis=0)
        Bin_summary=Bin_summary.T
        try:
            Bin_summary['Yield']=round(((Bin_summary[1]/Bin_summary['Total'])*100),2)
        except:
            pass
        Bin_summary=Bin_summary.T
        #Bin_summary['%'] = Bin_summary.apply(lambda row: "{0:.2f}%".format((row[value_field]/max(Bin_summary[value_field]))*100), axis = 1)
    #Hard_bin_sumary=Hard_bin_sumary.T
    Add_Data_Frame_to_Word(Bin_summary.T,document)
Binning_and_plotting(dataframe,'SoftBin','HardBin', 'Hardware')
Binning_and_plotting(dataframe,'HardBin','SoftBin', 'SoftWare')
Binning_and_plotting(dataframe,'HardBin','SoftBin', 'Site Wise','TestSiteNumber')
'''## Adding Hbin Summary to the Document
Hard_bin_sumary=pd.pivot_table(dataframe,values='SoftBin',index=['HardBin'],aggfunc=len,fill_value=0)
Hard_bin_sumary=Hard_bin_sumary.sort_values(by='SoftBin',ascending=False)
ax=Hard_bin_sumary.plot(kind='bar',x=Hard_bin_sumary.index,y='SoftBin',stacked=False,grid=False)
plt.ylabel('Count')
plt.title('Hardware Bin Pareto')
plt.savefig('Hbin_Pareto.png')
document.add_picture('Hbin_Pareto.png')
document.add_paragraph()
Hard_bin_sumary.loc['Total']= Hard_bin_sumary.sum(numeric_only=True, axis=0)
Hard_bin_sumary['%'] = Hard_bin_sumary.apply(lambda row: "{0:.2f}%".format((row.SoftBin/max(Hard_bin_sumary['SoftBin']))*100), axis = 1)
#Hard_bin_sumary=Hard_bin_sumary.T
Add_Data_Frame_to_Word(Hard_bin_sumary.T,document)'''





'''
table1=pd.pivot_table(dataframe,values='HardBin',index=['SoftBin'],aggfunc=len,fill_value=0)
table1=table1.sort_values(by='HardBin',ascending=False)
ax=table1.plot(kind='bar',x=table1.index,y='HardBin',stacked=False,grid=True)
table1.loc['Total']= table1.sum(numeric_only=True, axis=0)

## Adding Hbin Summary to the Document
table3=pd.pivot_table(dataframe,values='SoftBin',index=['HardBin'],aggfunc=len,fill_value=0)
table3=table3.sort_values(by='SoftBin',ascending=False)
table3.loc['Total']= table3.sum(numeric_only=True, axis=0)

table2=pd.pivot_table(dataframe,values='HardBin',index='TestSiteNumber',columns=['SoftBin'],aggfunc=len,fill_value=0)
table2.loc['Total']= table2.sum(numeric_only=True, axis=0)
table2=table2.sort_values(by='Total',ascending=False,axis=1)


Add_Data_Frame_to_Word(table1,document)
plt.savefig('Sbin.png')
document.add_picture('Sbin.png')
document.add_paragraph()
Add_Data_Frame_to_Word(table3,document)
Add_Data_Frame_to_Word(table2,document)'''
'''
p.add_run(' and some ')
p.add_run('italic.').italic = True
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')
document.add_page_break()
'''
document.save('demo4.docx')
print('Completed')