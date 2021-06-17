# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'Yiled_Report_Generation_form.ui'
#
# Created by: PyQt5 UI code generator 5.11.3
#
# WARNING! All changes made in this file will be lost!

from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import (QInputDialog,QFileDialog,QMessageBox,QWidget)
from docx import Document
from docx.shared import Inches,Mm
from docx.enum.style import WD_STYLE_TYPE
from pandas import DataFrame as DF
from pandas import *
import pandas as pd
import matplotlib.pyplot as plt
import os
class Ui_Yiled_Report_Generation_form(QtWidgets.QDialog):
    def __init__(self,parent=None):
        super(Ui_Yiled_Report_Generation_form,self).__init__()
        self.Parent_window=parent
    def setupUi(self, Yiled_Report_Generation_form):
        self.Yiled_report_window=Yiled_Report_Generation_form
        Yiled_Report_Generation_form.setObjectName("Yiled_Report_Generation_form")
        Yiled_Report_Generation_form.resize(381, 258)
        Yiled_Report_Generation_form.setMaximumSize(QtCore.QSize(381, 258))
        Yiled_Report_Generation_form.setWindowTitle('Yiled_Report_Generation_form')
        self.comboBox = QtWidgets.QComboBox(Yiled_Report_Generation_form)
        self.comboBox.setGeometry(QtCore.QRect(70, 10, 301, 21))
        self.comboBox.setObjectName("comboBox")
        self.groupBox = QtWidgets.QGroupBox(Yiled_Report_Generation_form)
        self.groupBox.setGeometry(QtCore.QRect(10, 40, 361, 211))
        self.groupBox.setTitle("")
        self.groupBox.setObjectName("groupBox")
        self.Site_Check_box = QtWidgets.QCheckBox(self.groupBox)
        self.Site_Check_box.setGeometry(QtCore.QRect(20, 80, 111, 17))
        self.Site_Check_box.setObjectName("Site_Check_box")
        self.Top_Failing_Test_Details = QtWidgets.QCheckBox(self.groupBox)
        self.Top_Failing_Test_Details.setGeometry(QtCore.QRect(20, 110, 131, 17))
        self.Top_Failing_Test_Details.setObjectName("Top_Failing_Test_Details")
        self.Histogram_for_Top_failing_Tests = QtWidgets.QCheckBox(self.groupBox)
        self.Histogram_for_Top_failing_Tests.setGeometry(QtCore.QRect(20, 140, 171, 17))
        self.Histogram_for_Top_failing_Tests.setObjectName("Histogram_for_Top_failing_Tests")
        self.HBin_Check_box = QtWidgets.QCheckBox(self.groupBox)
        self.HBin_Check_box.setGeometry(QtCore.QRect(20, 20, 111, 17))
        self.HBin_Check_box.setObjectName("checkBox")
        self.HBin_Check_box.setText('HardBin_Summary')
        self.checkBox_6 = QtWidgets.QCheckBox(self.groupBox)
        self.checkBox_6.setGeometry(QtCore.QRect(20, 170, 70, 17))
        self.checkBox_6.setObjectName("checkBox_6")
        self.SBin_Check_box = QtWidgets.QCheckBox(self.groupBox)
        self.SBin_Check_box.setGeometry(QtCore.QRect(20, 50, 111, 17))
        self.SBin_Check_box.setObjectName("SBin_Check_box")
        self.Ok_Button = QtWidgets.QPushButton(self.groupBox)
        self.Ok_Button.setGeometry(QtCore.QRect(220, 20, 75, 23))
        self.Ok_Button.setObjectName("Ok_Button")
        self.Ok_Button.clicked.connect(self.start_the_yiled_report_generation_process)
        self.Cancel_Button = QtWidgets.QPushButton(self.groupBox)
        self.Cancel_Button.setGeometry(QtCore.QRect(220, 60, 75, 23))
        self.Cancel_Button.setObjectName("Cancel_Button")
        self.Cancel_Button.clicked.connect(self.close_the_yield_report_window)
        self.label = QtWidgets.QLabel(Yiled_Report_Generation_form)
        self.label.setGeometry(QtCore.QRect(10, 10, 51, 16))
        self.label.setObjectName("label")

        self.Site_Check_box.setText( "SiteWise_Summary")
        self.Top_Failing_Test_Details.setText( "Top_Failing_Test_Details")
        self.Histogram_for_Top_failing_Tests.setText( "Histogram_for_Top_failing_Tests")
        #self.checkBox.setText( "HardBin_Summary")
        self.checkBox_6.setText( "CheckBox")
        self.SBin_Check_box.setText( "SoftBin_Summary")
        self.Ok_Button.setText( "OK")
        self.Cancel_Button.setText( "Cancel")
        self.label.setText( "File_Name:")        
    
    def close_the_yield_report_window(self):
        '''Closes the Yiled report window '''
        self.Yiled_report_window.close()
    def Get_the_loaded_file_details(self):
        '''Gets the loaded file data into this window'''
        if self.Parent_window!=None:
            for i in self.Parent_window.Loaded_Data_Files.keys():                
                self.comboBox.addItem(i)                
            if len(self.Parent_window.Loaded_Data_Files)>0: self.comboBox.setCurrentIndex(0)
        self.HBin_Check_box.setChecked(True)
        self.Site_Check_box.setChecked(True)
        self.Top_Failing_Test_Details.setChecked(True)
        self.Histogram_for_Top_failing_Tests.setChecked(True)      
        self.checkBox_6.setChecked(True)
        self.SBin_Check_box.setChecked(True)
    #Adding table to word doc function
    def Add_Data_Frame_to_Word(self,DataFrame,doc_name):
        df=DataFrame  # Assigning the Dataframe to df Variable
        doc_table = doc_name.add_table(df.shape[0]+1, df.shape[1]+1)
        
        # add the header rows.
        for j in range(df.shape[-1]):
            #print(df.shape)
            if j==0: 
                if df.columns.name!=None: doc_table.cell(0,j).text = df.columns.name
                if df.columns[j]!= None and df.columns[j]!='':
                    if type(df.columns[j]) is not str:
                        doc_table.cell(0,j+1).text = 'Bin '+str(df.columns[j])
                    else: 
                        doc_table.cell(0,j+1).text = df.columns[j]
            else: 
                if df.columns[j]!= None and df.columns[j]!='': 
                    if type(df.columns[j]) is not str:
                        doc_table.cell(0,j+1).text = 'Bin '+str(df.columns[j])
                    else: 
                        doc_table.cell(0,j+1).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            doc_table.cell(i+1,0).text = str(df.index[i])
            for j in range(df.shape[-1]):
                doc_table.cell(i+1,j+1).text = str(df.values[i,j])
        doc_table.style = 'TableGrid'
        doc_name.add_paragraph()
    def Binning_and_plotting(self,Dataframe_name,doc_name,value_field,index_column,Name_of_binning,columns_column=None):
        ## Adding Hbin Summary to the Document
        Bin_summary=pd.pivot_table(Dataframe_name,values=value_field,index=[index_column],columns=columns_column,aggfunc=len,fill_value=0)
        side_heading=doc_name.add_heading('Over All %s Binning:'%Name_of_binning, level=2)
        side_heading.bold = True
        plt.rcParams["figure.figsize"] = (7,3)
        if columns_column==None:
            Bin_summary=Bin_summary.sort_values(by=value_field,ascending=False)
            Bin_summary.plot(kind='bar',x=Bin_summary.index,y=value_field,stacked=False,grid=False)
                
        elif columns_column!=None:
            Bin_summary['Total']= Bin_summary.sum(numeric_only=True, axis=1)
            Bin_summary=Bin_summary.sort_values(by='Total',ascending=False)
            del Bin_summary['Total']
            Bin_summary.plot(kind='bar',x=Bin_summary.index,stacked=False,grid=False)
            #plt.rcParams["figure.figsize"] = (7,3)
        plt.ylabel('Count')
        plt.title(Name_of_binning+' Bin Pareto')
        plt.tight_layout()
        #plt.figure(figsize=(7,2))        
        #plt.show()
        plt.savefig(Name_of_binning+'.png')
        plt.close()       
        doc_name.add_picture(Name_of_binning+'.png')
        doc_name.add_paragraph()
        if columns_column==None:        
            Bin_summary.loc['Total']= Bin_summary.sum(numeric_only=True, axis=0)
            Bin_summary['%'] = Bin_summary.apply(lambda row: "{0:.2f}%".format((row[value_field]/max(Bin_summary[value_field]))*100), axis = 1)
        elif columns_column!=None: 
            Bin_summary.loc['Total']= Bin_summary.sum(numeric_only=True, axis=0)
            Bin_summary=Bin_summary.T
            try:
                Bin_summary['Yield']=round(((Bin_summary['1']/Bin_summary['Total'])*100),2)
            except:
                pass
            Bin_summary=Bin_summary.T
            #Bin_summary['%'] = Bin_summary.apply(lambda row: "{0:.2f}%".format((row[value_field]/max(Bin_summary[value_field]))*100), axis = 1)
        #Hard_bin_sumary=Hard_bin_sumary.T
        self.Add_Data_Frame_to_Word(Bin_summary.T,doc_name)

    def Add_Header_Fotter(self,doc_name):
        document=doc_name
        section=document.sections[0]
        #section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
        ## Adding Header and Fotter
        header=section.header
        header_table =header.add_table(rows=3, cols=4,width= Inches(9.0))
        header_table.cell(0,1).merge(header_table.cell(0,3))
        header_table.cell(0,0).merge(header_table.cell(2,0))
        header_table.cell(0,0).width=1.0
        #header_table.autofit==True
        header_table.cell(0,1).text='GOURISANKAR_TOOL'

        header_table.cell(1,1).text='Title'
        header_table.cell(1,1).width=Inches(0.01)
        header_table.cell(1,2).text='Test Result Report - Finger Print'
        header_table.cell(1,2).width=4.0
        header_table.cell(2,1).text='Doc No'
        header_table.cell(2,2).text='TBD'
        header_table.style = 'TableGrid'
        try:
            #import os
            icon_foldername=os.path.dirname(os.path.realpath(__file__))+"\\Icons\\"
            header_table.rows[0].cells[0].add_paragraph(0).add_run().add_picture(icon_foldername+'Logo.png',width=Inches(0.8))
        except:
            pass
        fotter=section.footer
        fotter.paragraphs[0].text='GOURISANKAR_TOOL'
        # Header and Fotter completed

    def start_the_yiled_report_generation_process(self):

        # Begining the documentation
        if self.Parent_window!=None and self.comboBox.count()>0:
            document=Document()
            self.Add_Header_Fotter(document)
            side_heading=document.add_heading('Introduction:', level=2)
            side_heading.bold = True
            document.add_paragraph('\t\t This report consists of yield and failure analysis for the lot DS5_DS5_T8C342-06F2_CP1_TS-HP9-014_DS5_CP  consist of 298 devices of wafer sort. ')
            #data=pd.read_csv('D:\\2360\\Scripting_data\\Python_trail\\GIT_HUB\\Tool\\STDF-READER-\\Input_to_clear\\Raw data\\W649X219_RT1_C633Y5F0301_7DBS5V_TS-HP9-014_20161206014854.csv')
            file_name=self.comboBox.currentText()
            dataframe=self.Parent_window.Loaded_Data_File_Raw_Data[self.comboBox.currentText()]
            dataframe['Count']=1
            #dataframe=DF(data)
            if 'MIR_Rec_Summary' in self.Parent_window.Loaded_Data_Files[file_name]:
                side_heading=document.add_heading('Product & test Program details:', level=2)
                side_heading.bold = True
                mir_table=DF(self.Parent_window.Loaded_Data_Files[file_name]['MIR_Rec_Summary'])
                Product_details_table=DF(mir_table,columns=['Setuptime', 'Strtuptime', 'Stat_Num', 'modcode', 'retestcode',
                'Burntime', 'Lotid', 'Parttype', 'TesterType', 'Jobname', 'Jobrev', 'Sublotid', 'Opername', 'ExSwtype', 'ExSwver',
                'TestCode', 'TestTemp', 'Pcktyp', 'FloorId', 'FabPId', 'OpFreq', 'TestSpecVerName', 'SetupId', 'DDrev', 'EnggId',
                'RomCode', 'NodName'])
                Product_details_table=Product_details_table[0:1].T
                Product_details_table.index.name=None
                Product_details_table.rename(columns = {0:" "}, inplace = True) 
                self.Add_Data_Frame_to_Word(Product_details_table,document)
            if self.HBin_Check_box.isChecked()==True:self.Binning_and_plotting(dataframe,document,'Count','HardBin', 'Hardware')
            if self.SBin_Check_box.isChecked()==True:self.Binning_and_plotting(dataframe,document,'Count','SoftBin', 'SoftWare')
            if self.Site_Check_box.isChecked()==True:self.Binning_and_plotting(dataframe,document,'Count','SoftBin', 'Site Wise','TestSiteNumber')
            if self.Top_Failing_Test_Details.isChecked()==True and 'TSR_Rec_Summary' in self.Parent_window.Loaded_Data_Files[file_name]:                
                Tsr=self.Parent_window.Loaded_Data_Files[file_name]['TSR_Rec_Summary']
                tsr_df=DF(Tsr)
                tsr_df['Test_number;Name']=tsr_df['TestNumber'].astype(str)+";"+tsr_df['Test_nam']
                table1=DF(tsr_df,columns=['Test_number;Name','Fail_Cnt','Test_type']).sort_values(by='Fail_Cnt',ascending=False)
                table1.set_index('Test_number;Name', inplace=True)
                table1.columns.name='TestNumber;Test_Name'
                #table1.sort_values(by='Fail_Cnt',ascending=False)
                table2=tsr_df.pivot_table(index=['TestNumber'],values='Fail_Cnt').sort_values(by='Fail_Cnt',ascending=False)
                table2.columns.name='Test Number'                
                plt.rcParams["figure.figsize"] = (7,3)
                table2.head(10).plot(kind='bar')
                #ax=Bin_summary.plot(kind='bar',x=Bin_summary.index,stacked=False,grid=False)
                plt.ylabel('Count')
                Name_of_binning='Top_Failing_Tests_Details'
                plt.title(Name_of_binning)
                plt.tight_layout()                
                #plt.show()
                plt.savefig(Name_of_binning+'.png')
                plt.close()
                side_heading=document.add_heading('Top 10 Failing Test Details:', level=2)
                side_heading.bold = True
                document.add_picture(Name_of_binning+'.png')
                self.Add_Data_Frame_to_Word(table1.head(10),document)
            if self.Histogram_for_Top_failing_Tests.isChecked()==True:
                for i in range(0,2):  # Need to have the count based
                    if table1['Test_type'][0]!='F':
                        data=self.Parent_window.Loaded_Data_File_Raw_Data[file_name][table1.index[i]]
                        sumary_table={'0':['TestNumber;Name'],'1' :[data.name]}
                        plt.rcParams["figure.figsize"] = (7,3)
                        data.plot(kind='hist')                        
                        if 'Test_Limit_Details' in self.Parent_window.Loaded_Data_Files[file_name]:
                            lo_lim=self.Parent_window.Loaded_Data_Files[file_name]['Test_Limit_Details'][table1.index[i]][0]
                            hi_lim=self.Parent_window.Loaded_Data_Files[file_name]['Test_Limit_Details'][table1.index[i]][1]
                            unit=self.Parent_window.Loaded_Data_Files[file_name]['Test_Limit_Details'][table1.index[i]][2]
                            Lims=[lo_lim,hi_lim]
                            colors=['b','r']
                            Labels=['Lo_Limit','Hi_Limit']
                            for li,c,lbl in zip(Lims,colors,Labels):
                                plt.axvline(x=li, c=c, label= lbl) if li!='' else ''
                                sumary_table['0'].append(lbl)
                                sumary_table['1'].append(li)
                            sumary_table['0'].append('Unit')
                            sumary_table['1'].append(unit)                        
                        plt.title(data.name)
                        plt.legend(loc='center left')#, bbox_to_anchor=(1, 0.5), borderaxespad=0., prop={'size': 10})
                        plt.title(data.name)
                        plt.tight_layout()
                        plt.savefig(str(data.name.split(";")[0])+'.png')
                        #plt.show()
                        plt.close()
                        summary_for_test=data.describe(include='all')
                        for i,j in zip(summary_for_test.index,summary_for_test.values):
                            sumary_table['0'].append(i)
                            sumary_table['1'].append(round(j,3))
                        if 'Test_Limit_Details' in self.Parent_window.Loaded_Data_Files[file_name]:
                            if lo_lim!='' and lo_lim!=None and hi_lim!='' and hi_lim!=None:                                
                                stddev=summary_for_test.values[2]
                                mean=summary_for_test.values[1]
                                if stddev!=0 and stddev!=None:
                                    Cp=round(((hi_lim -lo_lim)/(6*stddev)),3)
                                    CpK=round(min(((hi_lim-mean)/(3*stddev)),((mean-lo_lim)/(3*stddev))),3)
                                else:
                                    Cp='N/A'
                                    CpK='N/A'
                            sumary_table['0'].append('CP')
                            sumary_table['1'].append(Cp)
                            sumary_table['0'].append('CPK')
                            sumary_table['1'].append(CpK)
                        sumary_table=DF(sumary_table)                        
                        sumary_table.set_index('0', inplace=True)
                        sumary_table=sumary_table.T
                        sumary_table.set_index('TestNumber;Name', inplace=True)
                        sumary_table.columns.name='TestNumber;Name'
                        side_heading=document.add_heading('Histogram for %s'%data.name, level=2)
                        side_heading.bold = True
                        document.add_picture(data.name.split(";")[0]+'.png')
                        self.Add_Data_Frame_to_Word(sumary_table,document)
                        #plt.show()
            #Single header
            #import os
            filename=dataframe['File_Name'][0]
            filename=os.path.splitext(filename)[0]
            try:
                document.save(filename+'.docx')
            except:
                document.save(filename+'_.docx')
            print('Completed')
            self.close_the_yield_report_window()

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Yiled_Report_Generation_form = QtWidgets.QWidget()
    ui = Ui_Yiled_Report_Generation_form()
    ui.setupUi(Yiled_Report_Generation_form)
    Yiled_Report_Generation_form.show()
    sys.exit(app.exec_())

